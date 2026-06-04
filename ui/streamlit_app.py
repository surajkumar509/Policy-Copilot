import sys
import os
from dotenv import load_dotenv

# ✅ Setup path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
load_dotenv()

import random
import time
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
import io

import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# ✅ Imports
import agent.tools as tools
from rag.loader import load_documents
from rag.chunker import chunk_text
from rag.azure_embeddings import embed_text
from rag.shared_store import vector_db
from agent.agent import agent_run

# =========================
# ✅ SESSION STATE INIT
# =========================
if "loaded_sources" not in st.session_state:
    st.session_state.loaded_sources = set()

if "last_sync_time" not in st.session_state:
    st.session_state.last_sync_time = 0

if "load_error" not in st.session_state:
    st.session_state.load_error = None

if "initial_load_done" not in st.session_state:
    st.session_state.initial_load_done = False

if "initial_load_count" not in st.session_state:
    st.session_state.initial_load_count = 0

if "last_user_query" not in st.session_state:
    st.session_state.last_user_query = ""

if "last_response" not in st.session_state:
    st.session_state.last_response = ""

if "last_source" not in st.session_state:
    st.session_state.last_source = ""

if "last_sim_receipt" not in st.session_state:
    st.session_state.last_sim_receipt = None

if "simulated_email_history" not in st.session_state:
    st.session_state.simulated_email_history = []

if "download_choice" not in st.session_state:
    st.session_state.download_choice = "No"

if "download_format" not in st.session_state:
    st.session_state.download_format = "DOCX"

if "download_ready" not in st.session_state:
    st.session_state.download_ready = False

if "download_payload" not in st.session_state:
    st.session_state.download_payload = b""

if "download_filename" not in st.session_state:
    st.session_state.download_filename = ""

if "download_mime" not in st.session_state:
    st.session_state.download_mime = ""

if "download_label" not in st.session_state:
    st.session_state.download_label = ""

if "prepared_download_format" not in st.session_state:
    st.session_state.prepared_download_format = ""

if "user_query_text" not in st.session_state:
    st.session_state.user_query_text = ""

if "selected_country" not in st.session_state:
    st.session_state.selected_country = "All Countries"

COUNTRY_ALIASES = {
    "Canada": ["canada", "acnanda", "ca"],
    "USA": ["usa", "us", "united states", "america"],
    "India": ["india", "in"],
    "Japan": ["japan", "jp"],
    "Netherlands": ["netherlands", "neterland", "holland", "nl"],
    "UK": ["uk", "united kingdom", "britain", "gb"],
    "Germany": ["germany", "de"],
    "Singapore": ["singapore", "sg"],
    "Australia": ["australia", "au"],
}

COUNTRY_OPTIONS = [
    "Canada",
    "USA",
    "India",
    "Japan",
    "Netherlands",
    "UK",
    "Germany",
    "Singapore",
    "Australia",
]


# =========================
# ✅ PARALLEL EMBEDDING
# =========================
def embed_chunks_parallel(chunks):
    with ThreadPoolExecutor(max_workers=5) as executor:
        return list(executor.map(embed_text, chunks))


def parse_email_subject_and_body(response_text, user_query):
    lines = [line.strip() for line in response_text.splitlines() if line.strip()]
    if not lines:
        return "Policy Update", response_text

    first = lines[0]
    if first.lower().startswith("subject:"):
        subject = first.split(":", 1)[1].strip() or "Policy Update"
        body = "\n".join(response_text.splitlines()[1:]).strip()
        return subject, (body if body else response_text)

    short_query = user_query.strip()[:60] if user_query.strip() else "Policy Update"
    return f"Policy Follow-up: {short_query}", response_text


def build_download_filename(prefix, query):
    cleaned = "".join(
        ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in query.lower().strip()
    )
    cleaned = "_".join([part for part in cleaned.split("_") if part])
    if not cleaned:
        cleaned = "policy"
    return f"{prefix}_{cleaned[:40]}"


def build_checklist_pdf(query, response):
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 50
    line_height = 14

    def draw_line(text):
        nonlocal y
        if y < 50:
            pdf.showPage()
            y = height - 50
        pdf.drawString(40, y, text)
        y -= line_height

    pdf.setFont("Helvetica-Bold", 14)
    draw_line("Policy Checklist")
    y -= 4
    pdf.setFont("Helvetica-Bold", 11)
    draw_line("Query:")
    pdf.setFont("Helvetica", 10)
    for part in query.splitlines() or [query]:
        for chunk in [part[i : i + 95] for i in range(0, len(part), 95)] or [""]:
            draw_line(chunk)
    y -= 4
    pdf.setFont("Helvetica-Bold", 11)
    draw_line("Response:")
    pdf.setFont("Helvetica", 10)
    for part in response.splitlines() or [response]:
        for chunk in [part[i : i + 95] for i in range(0, len(part), 95)] or [""]:
            draw_line(chunk)

    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()


def build_checklist_docx(query, response):
    doc = Document()
    doc.add_heading("Policy Checklist", level=1)
    doc.add_heading("Query", level=2)
    doc.add_paragraph(query)
    doc.add_heading("Response", level=2)
    doc.add_paragraph(response)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_checklist_txt(query, response):
    content = (
        "Policy Checklist\n"
        "================\n\n"
        "Query:\n" + query + "\n\n" + "Response:\n" + response + "\n"
    )
    return content.encode("utf-8")


def format_last_sync_time(ts):
    if not ts:
        return "Not synced yet"
    dt = datetime.fromtimestamp(ts)
    return dt.strftime("%d %b %Y, %I:%M %p")


def extract_country_from_source(source):
    normalized = (source or "").replace("\\", "/")
    search_text = " ".join(
        normalized.lower().replace("/", " ").replace("_", " ").replace("-", " ").split()
    )

    for country, aliases in COUNTRY_ALIASES.items():
        for alias in aliases:
            alias_clean = alias.strip().lower()
            if alias_clean and alias_clean in search_text:
                return country

    parts = [part.strip() for part in normalized.split("/") if part.strip()]
    if len(parts) >= 2:
        country = parts[0].replace("_", " ").replace("-", " ").strip()
        return country.title() if country else "Uncategorized"
    return "Uncategorized"


def inject_enterprise_theme():
    st.markdown(
        """
        <style>
            @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;500;600;700&family=IBM+Plex+Mono:wght@400;600&display=swap');

            :root {
                --brand-navy: #0e1f3d;
                --brand-blue: #1556d6;
                --brand-cyan: #0aa6c8;
                --surface: #f4f7fb;
                --card: #ffffff;
                --border: #d9e2ef;
                --text: #12263a;
                --muted: #4b6178;
                --ok: #198754;
            }

            .stApp {
                background:
                    radial-gradient(circle at 12% 18%, rgba(10, 166, 200, 0.14), transparent 30%),
                    radial-gradient(circle at 85% 10%, rgba(21, 86, 214, 0.10), transparent 35%),
                    linear-gradient(180deg, #f8fbff 0%, var(--surface) 58%, #eef3f9 100%);
                color: var(--text);
                font-family: 'IBM Plex Sans', sans-serif;
            }

            #MainMenu {
                visibility: hidden;
            }

            [data-testid='stHeader'] {
                background: transparent !important;
                padding: 0.5rem 1rem !important;
                border-bottom: none !important;
                min-height: 40px !important;
            }

            [data-testid='stStatusWidget'],
            [data-testid='stDecoration'] {
                display: none;
            }

            [data-testid='stHeader'] button {
                display: inline-flex !important;
                visibility: visible !important;
                opacity: 1 !important;
                z-index: 1000;
                background: transparent !important;
                border: none !important;
                color: #1556d6 !important;
            }

            [data-testid='stHeader'] button:hover {
                opacity: 0.8 !important;
            }

            /* Hide top-right Deploy and overflow (three dots) controls */
            [data-testid='stAppDeployButton'],
            [data-testid='stAppDeployButton'] *,
            [data-testid*='Deploy'],
            [data-testid*='deploy'],
            [data-testid='stToolbar'] a,
            [data-testid='stToolbar'] button[aria-label*='Deploy'],
            [data-testid='stToolbar'] button[title*='Deploy'],
            [data-testid='stToolbar'] button[aria-label*='menu'],
            [data-testid='stToolbar'] button[aria-label*='Menu'],
            [data-testid='stToolbar'] button[title*='menu'],
            [data-testid='stToolbar'] button[title*='Menu'] {
                display: none !important;
                visibility: hidden !important;
                width: 0 !important;
                height: 0 !important;
                overflow: hidden !important;
            }

            [data-testid='stSidebar'] {
                display: flex !important;
                visibility: visible !important;
            }

            .main .block-container {
                padding-top: 1.5rem;
                padding-bottom: 2rem;
            }

            .hero-shell {
                background: linear-gradient(110deg, var(--brand-navy), #14356a 58%, #1b4a93 100%);
                color: #ffffff;
                border-radius: 18px;
                border: 1px solid rgba(255, 255, 255, 0.14);
                padding: 1.25rem 1.4rem;
                box-shadow: 0 16px 32px rgba(11, 32, 67, 0.24);
                margin-bottom: 1rem;
                animation: heroFade 0.45s ease-out;
            }

            .hero-title {
                font-size: 1.55rem;
                font-weight: 700;
                letter-spacing: 0.2px;
                margin-bottom: 0.25rem;
            }

            .hero-sub {
                color: rgba(255, 255, 255, 0.90);
                font-size: 0.95rem;
                margin: 0;
            }

            .kpi-grid {
                display: grid;
                grid-template-columns: repeat(3, minmax(120px, 1fr));
                gap: 0.8rem;
                margin: 0.85rem 0 1.1rem;
            }

            .kpi-card {
                background: var(--card);
                border: 1px solid var(--border);
                border-radius: 14px;
                padding: 0.8rem 0.95rem;
                box-shadow: 0 6px 18px rgba(18, 38, 58, 0.06);
            }

            .kpi-label {
                color: var(--muted);
                font-size: 0.78rem;
                text-transform: uppercase;
                letter-spacing: 0.5px;
                margin: 0;
            }

            .kpi-value {
                color: var(--text);
                font-size: 1.18rem;
                font-weight: 700;
                margin: 0.2rem 0 0;
            }

            .panel {
                background: var(--card);
                border: 1px solid var(--border);
                border-radius: 14px;
                padding: 1rem;
                box-shadow: 0 10px 24px rgba(18, 38, 58, 0.05);
                margin-bottom: 0.95rem;
            }

            div[data-testid='stForm'] {
                background: var(--card);
                border: 1px solid var(--border);
                border-radius: 14px;
                padding: 1rem 1rem 0.25rem;
                box-shadow: 0 10px 24px rgba(18, 38, 58, 0.05);
                margin-bottom: 0.95rem;
            }

            .query-chip {
                display: inline-block;
                background: rgba(21, 86, 214, 0.10);
                color: #103469;
                border: 1px solid rgba(21, 86, 214, 0.28);
                border-radius: 999px;
                font-size: 0.78rem;
                font-weight: 600;
                padding: 0.25rem 0.65rem;
                margin-bottom: 0.55rem;
            }

            .response-shell {
                background: #f8fbff;
                border: 1px solid #d6e3f5;
                border-left: 5px solid var(--brand-blue);
                border-radius: 12px;
                padding: 0.95rem 1rem;
                color: var(--text);
                line-height: 1.55;
            }

            [data-testid='stSidebar'] {
                background: linear-gradient(180deg, #0f2649 0%, #112f5a 58%, #12386f 100%);
                border-right: 1px solid rgba(255, 255, 255, 0.08);
            }

            [data-testid='stSidebar'] * {
                color: #eaf1ff;
            }

            [data-testid='stSidebar'] .stButton > button {
                border-radius: 10px;
                border: 1px solid rgba(255, 255, 255, 0.24);
                background: rgba(255, 255, 255, 0.08);
                color: #f4f8ff;
            }

            .stButton > button, .stDownloadButton > button, .stFormSubmitButton > button {
                border-radius: 10px;
                border: 1px solid #184aa3;
                background: linear-gradient(180deg, #1d5fd4 0%, #1550ba 100%);
                color: #ffffff;
                font-weight: 600;
            }

            .stButton > button:hover, .stDownloadButton > button:hover, .stFormSubmitButton > button:hover {
                transform: translateY(-1px);
                box-shadow: 0 8px 18px rgba(21, 86, 214, 0.22);
            }

            .stTextArea textarea, .stTextInput input {
                border-radius: 10px !important;
                border: 1px solid #bfd0e8 !important;
                background-color: #ffffff !important;
            }

            [data-testid='stSidebar'] .stTextInput input {
                color: #12263a !important;
                -webkit-text-fill-color: #12263a !important;
                caret-color: #12263a !important;
            }

            [data-testid='stSidebar'] .stTextInput input::placeholder {
                color: #6f8298 !important;
                opacity: 1;
            }

            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] > div {
                background: #ffffff !important;
                border: 1px solid #bfd0e8 !important;
                border-radius: 10px !important;
                color: #12263a !important;
            }

            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] input {
                color: #12263a !important;
                -webkit-text-fill-color: #12263a !important;
            }

            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] span,
            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] div,
            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] p {
                color: #12263a !important;
                -webkit-text-fill-color: #12263a !important;
            }

            div[data-baseweb='popover'] ul,
            div[data-baseweb='popover'] li,
            div[data-baseweb='popover'] span,
            div[data-baseweb='popover'] div {
                color: #12263a !important;
                -webkit-text-fill-color: #12263a !important;
            }

            [data-testid='stSidebar'] .stSelectbox [data-baseweb='select'] svg {
                fill: #35516f !important;
            }

            [data-testid='stSidebar'] [data-testid='stExpander'],
            [data-testid='stSidebar'] [data-testid='stExpander'] details,
            [data-testid='stSidebar'] [data-testid='stExpander'] details[open],
            [data-testid='stSidebar'] [data-testid='stExpander'] summary,
            [data-testid='stSidebar'] [data-testid='stExpander'] summary:hover,
            [data-testid='stSidebar'] [data-testid='stExpander'] div {
                background: transparent !important;
                background-color: transparent !important;
                color: #eaf1ff !important;
            }

            [data-testid='stSidebar'] [data-testid='stExpander'] details {
                border: 1px solid rgba(255, 255, 255, 0.18) !important;
                border-radius: 10px !important;
            }

            .cost-panel {
                border: 1px solid rgba(255, 255, 255, 0.15);
                border-radius: 12px;
                background: rgba(255, 255, 255, 0.06);
                padding: 0.75rem;
                margin-top: 0.6rem;
            }

            .cost-row {
                display: flex;
                justify-content: space-between;
                font-size: 0.85rem;
                margin: 0.2rem 0;
            }

            .cost-ok {
                color: #9ef1bc;
                font-weight: 700;
                margin-top: 0.35rem;
            }

            @keyframes heroFade {
                from { opacity: 0; transform: translateY(10px); }
                to { opacity: 1; transform: translateY(0); }
            }

            @media (max-width: 900px) {
                .kpi-grid {
                    grid-template-columns: 1fr;
                }
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


# =========================
# ✅ LOAD NEW DOCUMENTS
# =========================
def load_new_documents():
    try:
        docs = load_documents()
        added_count = 0
        st.session_state.load_error = None

        for doc in docs:

            # ✅ Skip already loaded
            if doc["source"] in st.session_state.loaded_sources:
                print(f"⏭️ Skipped: {doc['source']}")
                continue

            try:
                # ✅ STEP 1: Chunk text
                chunks = chunk_text(doc["text"])

                # ✅ STEP 3: LIMIT chunks (VERY IMPORTANT)
                MAX_CHUNKS = 80
                chunks = chunks[:MAX_CHUNKS]

                print(f"📊 {doc['source']} → {len(chunks)} chunks")

                # ✅ STEP 2: Embed (ONLY strings)
                vectors = embed_chunks_parallel(chunks)

                # ✅ Prepare records
                chunk_records = [
                    {"source": doc["source"], "text": chunk} for chunk in chunks
                ]

                # ✅ Store in vector DB
                vector_db.add(vectors, chunk_records)

                # ✅ Mark as loaded
                st.session_state.loaded_sources.add(doc["source"])
                added_count += 1

                print(f"✅ Indexed: {doc['source']}")

            except Exception as e:
                print(f"❌ Error processing {doc['source']}: {e}")

        st.session_state.last_sync_time = time.time()

        return added_count

    except Exception as e:
        st.session_state.load_error = str(e)
        print(f"❌ Load error: {e}")
        return 0


# =========================
# ✅ PAGE CONFIG
# =========================
st.set_page_config(
    page_title="Policy & SOP Copilot",
    layout="wide",
    initial_sidebar_state="expanded",
)
inject_enterprise_theme()

st.markdown(
    """
    <div class="hero-shell">
        <div class="hero-title">Compliance Checklist Copilot</div>
        <p class="hero-sub">Enterprise-grade assistant for policy Q&A, checklist automation, and communication workflows.</p>
    </div>
    """,
    unsafe_allow_html=True,
)

# One-time auto load: runs only once per session so reruns stay fast.
if not st.session_state.initial_load_done:
    with st.spinner("First-time setup: loading policy documents..."):
        st.session_state.initial_load_count = load_new_documents()
        st.session_state.initial_load_done = True

# =========================
# ✅ SIDEBAR
# =========================
with st.sidebar:
    st.header("Control Center")
    st.caption("Document sync, prompts, and usage insights")
    st.write(f"Loaded policies: {len(st.session_state.loaded_sources)}")

    # ✅ Manual Sync Button
    if st.button("Sync Knowledge Base", use_container_width=True):
        with st.spinner("Syncing policies..."):
            new_count = load_new_documents()

        if new_count > 0:
            st.success(f"✅ {new_count} new policies added")
        else:
            st.info("No new policies found")

    st.caption(f"Last sync: {format_last_sync_time(st.session_state.last_sync_time)}")

    st.subheader("Quick Prompts")
    quick_prompts = [
        "Show recent updates in policies",
        "Create a checklist for lta policy",
        "Draft an email for certification policy",
    ]
    for prompt in quick_prompts:
        if st.button(prompt, use_container_width=True):
            st.session_state.user_query_text = prompt
            st.session_state.quick_prompt_triggered = True
            st.rerun()

    st.subheader("Loaded Policies")

    if st.session_state.loaded_sources:
        policy_country_map = {
            policy: extract_country_from_source(policy)
            for policy in st.session_state.loaded_sources
        }
        available_countries = sorted(
            {
                country
                for country in policy_country_map.values()
                if country != "Uncategorized"
            }
        )
        country_options = ["All Countries"] + sorted(
            set(COUNTRY_OPTIONS).union(available_countries)
        )

        if st.session_state.selected_country not in country_options:
            st.session_state.selected_country = "All Countries"

        st.selectbox("Country", options=country_options, key="selected_country")
        policy_filter = st.text_input("Filter policies")
        visible_policies = sorted(st.session_state.loaded_sources)

        if st.session_state.selected_country != "All Countries":
            visible_policies = [
                p
                for p in visible_policies
                if policy_country_map.get(p) == st.session_state.selected_country
            ]

        if policy_filter.strip():
            visible_policies = [
                p
                for p in visible_policies
                if policy_filter.lower().strip() in p.lower()
            ]
        with st.expander(f"View list ({len(visible_policies)})", expanded=False):
            for i, policy in enumerate(visible_policies, 1):
                st.write(f"{i}. {policy}")
    else:
        st.write("No policies loaded")

    if st.session_state.load_error:
        st.error(st.session_state.load_error)

    st.markdown(
        f"""
        <div class="cost-panel">
            <div style="font-weight:700; margin-bottom: 0.35rem;">Token Cost Optimization</div>
            <div class="cost-row"><span>Total Queries</span><span>{tools.TOTAL_QUERIES}</span></div>
            <div class="cost-row"><span>Cache Hits</span><span>{tools.CACHE_HITS}</span></div>
            <div class="cost-row"><span>API Calls</span><span>{tools.API_CALLS}</span></div>
            <div class="cost-ok">Savings: {tools.get_cost_savings()}%</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.markdown(
    f"""
    <div class="kpi-grid">
        <div class="kpi-card">
            <p class="kpi-label">Policies Indexed</p>
            <p class="kpi-value">{len(st.session_state.loaded_sources)}</p>
        </div>
        <div class="kpi-card">
            <p class="kpi-label">Initial Session Load</p>
            <p class="kpi-value">{st.session_state.initial_load_count}</p>
        </div>
        <div class="kpi-card">
            <p class="kpi-label">Last Sync</p>
            <p class="kpi-value" style="font-size:0.95rem;">{format_last_sync_time(st.session_state.last_sync_time)}</p>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# =========================
# ✅ USER INPUT FORM
# =========================
with st.form("query_form", clear_on_submit=False):
    st.markdown(
        '<div class="query-chip">Policy Query Assistant</div>',
        unsafe_allow_html=True,
    )
    st.text_area(
        "Ask a policy question",
        key="user_query_text",
        height=130,
        placeholder="Type a policy question, checklist request, or an email drafting task...",
    )
    submitted = st.form_submit_button("Run Policy Analysis", use_container_width=True)

# =========================
# ✅ QUERY HANDLING
# =========================
quick_prompt_triggered = st.session_state.get("quick_prompt_triggered", False)
if submitted or quick_prompt_triggered:
    if quick_prompt_triggered:
        st.session_state.quick_prompt_triggered = False
    user_query = st.session_state.user_query_text
    if user_query.strip():
        # # ✅ LAZY LOAD (ONLY FIRST TIME)
        # if not st.session_state.initial_load_done:
        #     with st.spinner("🔄 Loading policies (one-time setup)..."):
        #         load_new_documents()
        #         st.session_state.initial_load_done = True

        # ✅ RUN AGENT
        with st.spinner("Thinking..."):
            response, source = agent_run(
                user_query,
                country_filter=st.session_state.get(
                    "selected_country", "All Countries"
                ),
            )

        # ✅ VARIATION LOGIC
        if source != "API_CALL":
            if random.random() < 0.2:
                response = tools.rephrase_response(response)
            else:
                response = tools.vary_response(response)

        st.session_state.last_user_query = user_query
        st.session_state.last_response = response
        st.session_state.last_source = source
        st.session_state.download_choice = "No"
        st.session_state.download_format = "DOCX"
        st.session_state.download_ready = False
        st.session_state.download_payload = b""
        st.session_state.download_filename = ""
        st.session_state.download_mime = ""
        st.session_state.download_label = ""
        st.session_state.prepared_download_format = ""

if st.session_state.last_response:
    user_query = st.session_state.last_user_query
    response = st.session_state.last_response
    source = st.session_state.last_source
    dot = "☑️" if source != "API_CALL" else "✅"

    formatted = response.replace("\n", "<br>")

    st.markdown(
        f"""
        <div class="panel">
            <div class="query-chip">Latest Query</div>
            <div style="font-size:0.96rem; color:#20374f; margin-bottom:0.65rem;">{user_query}</div>
            <div style="margin: 0.2rem 0 0.65rem; font-weight:700;">{dot} Response</div>
            <div class="response-shell">
                {formatted}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    normalized_query = tools.normalize_query(user_query)
    original_query_lower = user_query.lower()

    # ── Intent detection mirrors agent.py priority: email > checklist > answer ──
    is_email_intent = any(w in original_query_lower for w in ("draft", "email", "mail"))
    is_checklist_intent = (not is_email_intent) and ("checklist" in normalized_query)

    if is_checklist_intent:
        st.markdown("### Download Checklist")
        st.radio(
            "Would you like to download this checklist?",
            options=["No", "Yes"],
            key="download_choice",
            horizontal=True,
        )

        if st.session_state.download_choice == "Yes":
            st.radio(
                "Which format do you want?",
                options=["DOCX", "PDF", "TXT"],
                key="download_format",
                horizontal=True,
            )

            if st.button("Run Download Automation", use_container_width=True):
                with st.status("Automation running...", expanded=True) as status:
                    st.write("1. Reading checklist response")
                    st.write(f"2. Preparing {st.session_state.download_format} file")

                    base_name = build_download_filename("checklist", user_query)
                    selected_format = st.session_state.download_format

                    if selected_format == "DOCX":
                        payload = build_checklist_docx(user_query, response)
                        ext = ".docx"
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    elif selected_format == "PDF":
                        payload = build_checklist_pdf(user_query, response)
                        ext = ".pdf"
                        mime = "application/pdf"
                    else:
                        payload = build_checklist_txt(user_query, response)
                        ext = ".txt"
                        mime = "text/plain"

                    st.session_state.download_payload = payload
                    st.session_state.download_filename = f"{base_name}{ext}"
                    st.session_state.download_mime = mime
                    st.session_state.download_label = (
                        f"Download Checklist ({selected_format})"
                    )
                    st.session_state.download_ready = True
                    st.session_state.prepared_download_format = selected_format

                    st.write("3. File prepared successfully")
                    status.update(label="Automation completed", state="complete")

            if (
                st.session_state.download_ready
                and st.session_state.prepared_download_format
                == st.session_state.download_format
            ):
                st.success(
                    f"Automation done. Ready to download {st.session_state.download_filename}"
                )
                st.download_button(
                    label=st.session_state.download_label,
                    data=st.session_state.download_payload,
                    file_name=st.session_state.download_filename,
                    mime=st.session_state.download_mime,
                )

    if is_email_intent:
        st.markdown("### Email Simulation (No Real Send)")
        suggested_subject, suggested_body = parse_email_subject_and_body(
            response, user_query
        )

        with st.form("email_simulation_form"):
            to_raw = st.text_input("To", value="manager@company.com")
            cc_raw = st.text_input("CC (optional)", value="")
            subject = st.text_input("Subject", value=suggested_subject)
            body = st.text_area("Body", value=suggested_body, height=180)
            simulate_submit = st.form_submit_button("Simulate Send")

        if simulate_submit:
            recipients = [
                x.strip() for x in (to_raw + "," + cc_raw).split(",") if x.strip()
            ]
            st.session_state.last_sim_receipt = tools.simulate_email_send(
                recipients, subject, body
            )
            if st.session_state.last_sim_receipt["ok"]:
                st.session_state.simulated_email_history.append(
                    st.session_state.last_sim_receipt["details"]
                )

        if st.session_state.last_sim_receipt:
            receipt = st.session_state.last_sim_receipt
            if receipt["ok"]:
                detail = receipt["details"]
                st.success(
                    f"Mail send simulated successfully for {detail['recipient_count']} recipient(s). "
                    f"Reference: {detail['message_id']}"
                )
            else:
                st.error(receipt["message"])

        if st.session_state.simulated_email_history:
            with st.expander("View Simulation History"):
                for idx, item in enumerate(
                    reversed(st.session_state.simulated_email_history), 1
                ):
                    st.write(
                        f"{idx}. {item['timestamp']} | {item['message_id']} | {item['recipient_count']} recipient(s)"
                    )
elif submitted:
    st.warning("Please enter a query.")
